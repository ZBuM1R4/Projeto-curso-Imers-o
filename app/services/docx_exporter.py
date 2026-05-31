from datetime import datetime
from io import BytesIO
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# =========================================================
# 🎨 IDENTIDADE VISUAL
# =========================================================
COLOR_DARK = "0B1220"
COLOR_DARK_2 = "111827"
COLOR_BLUE = "4F7CFF"
COLOR_PURPLE = "7C3AED"
COLOR_TEXT = "111827"
COLOR_MUTED = "64748B"
COLOR_LIGHT_BG = "F8FAFC"
COLOR_BORDER = "E2E8F0"
COLOR_GREEN = "10B981"
COLOR_YELLOW = "F59E0B"
COLOR_RED = "EF4444"


# =========================================================
# 🧩 HELPERS DOCX
# =========================================================
def set_cell_background(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_border(cell, color: str = COLOR_BORDER, size: str = "8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_style(run, size=10.5, bold=False, color=COLOR_TEXT):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(
    document,
    text: str,
    size: float = 10.5,
    bold: bool = False,
    color: str = COLOR_TEXT,
    alignment=None,
    space_after: int = 6,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment if alignment else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(space_after)

    run = paragraph.add_run(text)
    set_run_style(run, size=size, bold=bold, color=color)

    return paragraph


def add_cell_paragraph(
    cell,
    text: str,
    size: float = 10,
    bold: bool = False,
    color: str = COLOR_TEXT,
    alignment=None,
    space_after: int = 4,
):
    paragraph = cell.add_paragraph()
    paragraph.alignment = alignment if alignment else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(space_after)

    run = paragraph.add_run(text)
    set_run_style(run, size=size, bold=bold, color=color)

    return paragraph


def clear_cell(cell):
    cell.text = ""


def add_section_title(document, title: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(8)

    run = paragraph.add_run(title)
    set_run_style(run, size=15, bold=True, color=COLOR_DARK)

    return paragraph


def add_small_label(document, label: str, value: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)

    label_run = paragraph.add_run(f"{label}: ")
    set_run_style(label_run, size=9.5, bold=True, color=COLOR_MUTED)

    value_run = paragraph.add_run(str(value))
    set_run_style(value_run, size=9.5, bold=False, color=COLOR_TEXT)

    return paragraph


def format_seconds(seconds) -> str:
    try:
        seconds = float(seconds)
    except Exception:
        return "0s"

    minutes = int(seconds // 60)
    remaining = int(seconds % 60)

    if minutes > 0:
        return f"{minutes}min {remaining}s"

    return f"{round(seconds, 2)}s"


def filter_attention_points(points: list[str]) -> list[str]:
    filtered = []

    for point in points:
        point_lower = point.lower()

        if "uso frequente" in point_lower:
            continue

        if "vício" in point_lower or "vicio" in point_lower:
            continue

        if "vícios" in point_lower or "vicios" in point_lower:
            continue

        filtered.append(point)

    return filtered


def filter_recurring_terms(report: dict) -> dict:
    repetitions = report.get("repeticoes", {})
    return repetitions.get("termos_recorrentes", {}) or {}


# =========================================================
# 📊 CHARTS
# =========================================================
def generate_bar_chart(
    title: str,
    labels: list[str],
    values: list[int],
    color: str = f"#{COLOR_BLUE}",
) -> BytesIO | None:
    if not labels or not values:
        return None

    fig, ax = plt.subplots(figsize=(7.6, 3.4), dpi=160)

    ax.bar(labels, values, color=color)
    ax.set_title(title, fontsize=13, fontweight="bold", color=f"#{COLOR_DARK}")
    ax.set_ylabel("Ocorrências", color=f"#{COLOR_MUTED}")
    ax.tick_params(axis="x", labelrotation=20)
    ax.tick_params(axis="y", colors=f"#{COLOR_MUTED}")

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(f"#{COLOR_BORDER}")
    ax.spines["bottom"].set_color(f"#{COLOR_BORDER}")

    ax.grid(axis="y", alpha=0.22)

    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    buffer.seek(0)
    plt.close(fig)

    return buffer


def generate_pie_chart(
    title: str,
    labels: list[str],
    values: list[float],
) -> BytesIO | None:
    if not labels or not values or sum(values) == 0:
        return None

    colors = [f"#{COLOR_BLUE}", f"#{COLOR_PURPLE}"]

    fig, ax = plt.subplots(figsize=(5.8, 3.6), dpi=160)

    ax.pie(
        values,
        labels=labels,
        autopct="%1.1f%%",
        startangle=90,
        colors=colors,
        textprops={"fontsize": 9, "color": f"#{COLOR_TEXT}"},
    )
    ax.set_title(title, fontsize=13, fontweight="bold", color=f"#{COLOR_DARK}")

    fig.patch.set_facecolor("white")

    buffer = BytesIO()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    buffer.seek(0)
    plt.close(fig)

    return buffer


# =========================================================
# 🧾 REPORT BLOCKS
# =========================================================
def add_cover(document: Document, report: dict, video_name: str):
    score_data = report.get("score_comunicacao", {})
    pausas = report.get("pausas", {})

    created_at = datetime.now().strftime("%d/%m/%Y")
    score = score_data.get("score", 0)
    classification = score_data.get("classificacao", "N/A")

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    left_cell.width = Inches(2.1)
    right_cell.width = Inches(4.8)

    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_cell_background(left_cell, COLOR_DARK)
    set_cell_border(left_cell, COLOR_DARK)
    set_cell_border(right_cell, COLOR_BORDER)

    clear_cell(left_cell)
    clear_cell(right_cell)

    add_cell_paragraph(
        left_cell,
        "ANÁLISE DE\nCOMUNICAÇÃO",
        size=15,
        bold=True,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    add_cell_paragraph(
        left_cell,
        "Relatório de Oratória",
        size=9,
        color="CBD5E1",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=22,
    )

    add_cell_paragraph(
        left_cell,
        "DATA DA ANÁLISE",
        size=8,
        bold=True,
        color="94A3B8",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        left_cell,
        created_at,
        size=10,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )

    add_cell_paragraph(
        left_cell,
        "ARQUIVO",
        size=8,
        bold=True,
        color="94A3B8",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        left_cell,
        Path(video_name).name,
        size=8.5,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )

    add_cell_paragraph(
        left_cell,
        "DURAÇÃO",
        size=8,
        bold=True,
        color="94A3B8",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        left_cell,
        format_seconds(pausas.get("duracao_total", 0)),
        size=10,
        color="FFFFFF",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        right_cell,
        "Relatório de Análise de Comunicação",
        size=20,
        bold=True,
        color=COLOR_DARK,
        space_after=8,
    )

    add_cell_paragraph(
        right_cell,
        "Avaliação gerada a partir da transcrição, métricas de fala e análise por inteligência artificial.",
        size=10.5,
        color=COLOR_MUTED,
        space_after=18,
    )

    score_table = right_cell.add_table(rows=1, cols=2)
    score_table.autofit = False

    score_cell = score_table.rows[0].cells[0]
    classification_cell = score_table.rows[0].cells[1]

    set_cell_background(score_cell, COLOR_LIGHT_BG)
    set_cell_background(classification_cell, COLOR_LIGHT_BG)
    set_cell_border(score_cell, COLOR_BORDER)
    set_cell_border(classification_cell, COLOR_BORDER)

    clear_cell(score_cell)
    clear_cell(classification_cell)

    add_cell_paragraph(
        score_cell,
        "SCORE GERAL",
        size=8,
        bold=True,
        color=COLOR_MUTED,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        score_cell,
        f"{score}/100",
        size=24,
        bold=True,
        color=COLOR_BLUE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        classification_cell,
        "CLASSIFICAÇÃO",
        size=8,
        bold=True,
        color=COLOR_MUTED,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        classification_cell,
        classification,
        size=14,
        bold=True,
        color=COLOR_PURPLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_cell_paragraph(
        right_cell,
        score_data.get("comentario", ""),
        size=10.5,
        color=COLOR_TEXT,
        space_after=6,
    )

    document.add_page_break()


def add_metrics_overview(document: Document, report: dict):
    score_data = report.get("score_comunicacao", {})
    pausas = report.get("pausas", {})
    recurring_terms = filter_recurring_terms(report)

    add_section_title(document, "Visão geral dos resultados")

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    metrics = [
        ("Score", f"{score_data.get('score', 0)}/100", COLOR_BLUE),
        ("Pausas longas", pausas.get("quantidade_pausas_longas", 0), COLOR_YELLOW),
        ("Silêncio", format_seconds(pausas.get("tempo_total_silencio", 0)), COLOR_PURPLE),
        ("Termos recorrentes", len(recurring_terms), COLOR_GREEN),
    ]

    for index, (label, value, color) in enumerate(metrics):
        cell = table.rows[0].cells[index]
        clear_cell(cell)
        set_cell_background(cell, COLOR_LIGHT_BG)
        set_cell_border(cell, COLOR_BORDER)

        add_cell_paragraph(
            cell,
            label.upper(),
            size=8,
            bold=True,
            color=COLOR_MUTED,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        add_cell_paragraph(
            cell,
            str(value),
            size=16,
            bold=True,
            color=color,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )


def add_executive_summary(document: Document, report: dict):
    add_section_title(document, "Resumo executivo")

    ai_analysis = report.get("analise_global_ia", {})
    score_data = report.get("score_comunicacao", {})

    if ai_analysis.get("analise"):
        add_paragraph(
            document,
            ai_analysis.get("analise", ""),
            size=10.5,
            color=COLOR_TEXT,
            space_after=8,
        )
    else:
        add_paragraph(
            document,
            score_data.get(
                "comentario",
                "Resumo indisponível para esta análise.",
            ),
            size=10.5,
            color=COLOR_TEXT,
            space_after=8,
        )


def add_attention_points(document: Document, report: dict):
    add_section_title(document, "Pontos de atenção")

    points = filter_attention_points(report.get("pontos_atencao", []))

    if not points:
        add_paragraph(
            document,
            "Nenhum ponto crítico identificado nesta análise.",
            color=COLOR_MUTED,
        )
        return

    for point in points:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(point)
        set_run_style(run, size=10.5, color=COLOR_TEXT)


def add_pauses_section(document: Document, report: dict):
    add_section_title(document, "Pausas e ritmo")

    pausas = report.get("pausas", {})

    add_small_label(document, "Duração total", format_seconds(pausas.get("duracao_total", 0)))
    add_small_label(document, "Pausas longas", pausas.get("quantidade_pausas_longas", 0))
    add_small_label(document, "Tempo total de silêncio", format_seconds(pausas.get("tempo_total_silencio", 0)))

    pausas_longas = pausas.get("pausas_longas", [])

    if pausas_longas:
        add_paragraph(document, "Pausas longas detectadas:", bold=True)

        for pausa in pausas_longas:
            document.add_paragraph(
                f"{round(pausa['start'], 2)}s → "
                f"{round(pausa['end'], 2)}s "
                f"({round(pausa['duration'], 2)}s)",
                style="List Bullet",
            )
    else:
        add_paragraph(document, "Nenhuma pausa longa detectada.", color=COLOR_MUTED)

    tempo_total = pausas.get("duracao_total", 0)
    tempo_silencio = pausas.get("tempo_total_silencio", 0)
    tempo_fala = max(tempo_total - tempo_silencio, 0)

    pie_chart = generate_pie_chart(
        title="Distribuição entre fala e silêncio",
        labels=["Fala", "Silêncio"],
        values=[tempo_fala, tempo_silencio],
    )

    if pie_chart:
        document.add_picture(pie_chart, width=Inches(4.9))


def add_repetitions_section(document: Document, report: dict):
    add_section_title(document, "Repetições e termos recorrentes")

    repetitions = report.get("repeticoes", {})
    sequential = repetitions.get("sequenciais", [])
    recurring = filter_recurring_terms(report)

    if sequential:
        add_paragraph(document, "Repetições em sequência:", bold=True)

        for item in sequential:
            document.add_paragraph(
                f"Termo repetido: {item.get('termo', '')}",
                style="List Bullet",
            )
    else:
        add_paragraph(document, "Nenhuma repetição em sequência.", color=COLOR_MUTED)

    if recurring:
        add_paragraph(document, "Termos mais recorrentes:", bold=True)

        for term, quantity in recurring.items():
            document.add_paragraph(f"{term}: {quantity}", style="List Bullet")

        chart = generate_bar_chart(
            title="Termos mais recorrentes",
            labels=list(recurring.keys()),
            values=list(recurring.values()),
            color=f"#{COLOR_PURPLE}",
        )

        if chart:
            document.add_picture(chart, width=Inches(5.8))
    else:
        add_paragraph(document, "Nenhum termo recorrente relevante.", color=COLOR_MUTED)


def add_transcription_section(document: Document, report: dict):
    document.add_page_break()

    add_section_title(document, "Transcrição")

    transcription = report.get("transcricao", "")

    if transcription:
        add_paragraph(document, transcription, size=10, color=COLOR_TEXT)
    else:
        add_paragraph(document, "Transcrição indisponível.", color=COLOR_MUTED)


def add_recommendations_section(document: Document):
    add_section_title(document, "Recomendação geral")

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.rows[0].cells[0]
    clear_cell(cell)
    set_cell_background(cell, COLOR_LIGHT_BG)
    set_cell_border(cell, COLOR_BORDER)

    add_cell_paragraph(
        cell,
        "Continue praticando com foco em clareza, organização da mensagem, ritmo e pausas estratégicas. "
        "Use este relatório como referência para acompanhar sua evolução nas próximas apresentações.",
        size=10.5,
        color=COLOR_TEXT,
    )


def add_human_review_section(document: Document):
    add_section_title(document, "Observações para revisão humana")

    add_paragraph(
        document,
        "Espaço reservado para comentários, correções e observações do revisor.",
        color=COLOR_MUTED,
    )

    for _ in range(4):
        add_paragraph(document, "____________________________________________________________________", color=COLOR_BORDER)


# =========================================================
# 🧾 MAIN
# =========================================================
def build_docx_report(report: dict, video_name: str = "Vídeo analisado") -> bytes:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    add_cover(document, report, video_name)
    add_metrics_overview(document, report)
    add_executive_summary(document, report)
    add_attention_points(document, report)
    add_pauses_section(document, report)
    add_repetitions_section(document, report)
    add_recommendations_section(document)
    add_transcription_section(document, report)
    add_human_review_section(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()